from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

agreement = [ 
        {
            'status': 'True',
            'comments': [
                {
                'clause': '',
                'modified': '',
                'original': ''
                }
            ],
            'name': 'Виноградов Рудольф Наумович',
            'user_id': 0,
            'role_name': 'юрист',
            'department': 'Юридический отдел'
        },
        {
            'status': 'True',
            'comments': [
                {
                'clause': '',
                'modified': '',
                'original': ''
                }
            ],
            'name': 'Ефимова Аделия Игоревна',
            'user_id': 1,
            'role_name': 'бухгалтер',
            'department': 'Бухгалтерия'
        },
        {
            'status': 'False',
            'comments': [
                {
                'clause': '2.1.',
                'modified': 'Срок окончания аренды: "10" июля 2022 г',
                'original': 'Срок окончания аренды: "01" июля 2022 г'
                }
            ],
            'name': 'Дубов Илья Елисеевич',
            'user_id': 2,
            'role_name': 'аренда',
            'department': 'Отдел аренды'
        },
        {
            'status': 'True',
            'comments': [
                {
                'clause': '',
                'modified': '',
                'original': ''
                }
            ],
            'name': 'Дубов Илья Елисеевич',
            'user_id': 2,
            'role_name': 'аренда',
            'department': 'Отдел аренды'
        },
        {
            'status': 'True',
            'comments': [
                {
                'clause': '',
                'modified': '',
                'original': ''
                }
            ],
            'name': 'Зайцев Илларион Денисович',
            'user_id': 3,
            'role_name': 'директор',
            'department': 'Директор'
        },
      ]

document = Document()

style = document.styles.add_style('UserHead1', WD_STYLE_TYPE.PARAGRAPH)
style.font.size = Pt(18)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

heading = document.add_paragraph('Лист согласования', style='UserHead1')
document.add_paragraph('Наименование договора: ______')
document.add_paragraph('Описание договора: ______')
document.add_paragraph('Контрагент: ______')
document.add_paragraph('Ответственный за согласование: Виноградов Рудольф Наумович')

r = len(agreement) + 1
table = document.add_table(rows=r, cols=7)
table.style = 'Table Grid'
cells0 = table.rows[0].cells
cells0[0].text = 'No п/п'
cells0[1].text = 'Наименование подразделения'
cells0[2].text = 'ФИО'
cells0[3].text = 'Комментируемый пункт'
cells0[4].text = 'Измененный пункт'
cells0[5].text = 'Дата согласования'
cells0[6].text = 'Результат согласования'

for x in range(1, r):
    cells = table.rows[x].cells
    cells[0].text = str(x)
    cells[1].text = str(agreement[x - 1].get('department'))
    cells[2].text = str(agreement[x - 1].get('name'))
    cells[3].text = str(agreement[x - 1].get('comments')[0].get('original'))
    cells[4].text = str(agreement[x - 1].get('comments')[0].get('modified'))
    cells[5].text = '21.04.2022'
    cells[6].text = 'Не согласовано' if agreement[x - 1].get('status') == 'False' else 'Согласовано'

document.add_paragraph('')
document.add_paragraph('Срок договора: _____')
document.save('Лист согласования.docx')
