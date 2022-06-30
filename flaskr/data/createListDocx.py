from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import aliased

from flaskr import db
from flaskr.models import Contract, CompaniesContract, Company, User, ContractVersionsUsers, Role, ContractVersion


def create_list_docx(c_id):
    userShouldAgree = aliased(User)
    roleShouldAgree = aliased(Role)
    userShouldSign = aliased(User)
    roleShouldSign = aliased(Role)

    cc = db.session.query(CompaniesContract.title,
                          Company.name.label('company_name'),
                          userShouldAgree.name.label('should_agreed_by_name'),
                          roleShouldAgree.name.label('should_agreed_by_role'),
                          Contract.agreed_at,
                          userShouldSign.name.label('should_signed_by_name'),
                          roleShouldSign.name.label('should_signed_by_role'),
                          Contract.signed_at)\
        .select_from(Contract)\
        .join(CompaniesContract)\
        .join(Company, Company.id == CompaniesContract.recipient_id)\
        .join(userShouldAgree, userShouldAgree.id == Contract.should_agreed_by)\
        .join(roleShouldAgree, roleShouldAgree.id == userShouldAgree.role_id)\
        .join(userShouldSign, userShouldSign.id == Contract.should_signed_by)\
        .join(roleShouldSign, roleShouldSign.id == userShouldSign.role_id)\
        .filter(Contract.id == c_id)\
        .first()

    agreements = db.session.query(User.name,
                                  Role.name.label('role_name'),
                                  ContractVersionsUsers.status_added_at)\
        .select_from(Contract)\
        .join(ContractVersion)\
        .join(ContractVersionsUsers)\
        .join(User, User.id == ContractVersionsUsers.user_id)\
        .join(Role)\
        .filter(Contract.id == c_id, ContractVersionsUsers.status == True)\
        .all()

    document = Document()

    style = document.styles.add_style('UserHead1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(18)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading = document.add_paragraph('Лист согласования', style='UserHead1')
    document.add_paragraph('Наименование договора: ' + cc.title)
    document.add_paragraph('Контрагент: ' + cc.company_name)
    document.add_paragraph('Ответственный за согласование: ' + cc.should_agreed_by_name)
    document.add_paragraph('Дата согласования: ' + cc.signed_at.strftime("%H:%M %d.%m.%Y"))

    r = len(agreements) + 3
    table = document.add_table(rows=r, cols=5)
    table.style = 'Table Grid'
    cells0 = table.rows[0].cells
    cells0[0].text = 'No п/п'
    cells0[1].text = 'Наименование подразделения'
    cells0[2].text = 'ФИО'
    cells0[3].text = 'Дата согласования'
    cells0[4].text = 'Результат согласования'

    for x in range(1, r-2):
        cells = table.rows[x].cells
        cells[0].text = str(x)
        cells[1].text = str(agreements[x - 1]['role_name'])
        cells[2].text = str(agreements[x - 1]['name'])
        cells[3].text = str(agreements[x - 1]['status_added_at'].strftime("%H:%M %d.%m.%Y"))
        cells[4].text = 'Согласовано'

    cells = table.rows[r-2].cells
    cells[0].text = str(r-2)
    cells[1].text = str(cc.should_agreed_by_role)
    cells[2].text = str(cc.should_agreed_by_name)
    cells[3].text = str(cc.agreed_at.strftime("%H:%M %d.%m.%Y"))
    cells[4].text = 'Завизировано'

    cells = table.rows[r - 1].cells
    cells[0].text = str(r-1)
    cells[1].text = str(cc.should_signed_by_role)
    cells[2].text = str(cc.should_signed_by_name)
    cells[3].text = str(cc.signed_at.strftime("%H:%M %d.%m.%Y"))
    cells[4].text = 'Подписано'

    document.add_paragraph('')
    document.save('files/Лист согласования.docx')
