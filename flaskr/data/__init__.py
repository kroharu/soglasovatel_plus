import datetime
import os
import re
import sys
import subprocess
import smtplib
from email.mime.text import MIMEText
from uuid import uuid4
from flask import Flask, render_template, request, jsonify, send_from_directory, flash
from werkzeug.utils import secure_filename
from werkzeug.datastructures import  FileStorage

import docx

from flaskr.models import User, Contract, Company, Role, ContractVersion, ContractVersionsUsers, UsersComments, \
    CompaniesContract, ContractsTypesUsers, ContractsType
from flaskr import db

app = Flask(__name__)

app.config['FOLDER_PDF'] = "./"

MAIL_SERVER = 'smtp.gmail.com'
MAIL_PORT = 465

MAIL_USERNAME = 'soglasovatel.plus.team2@gmail.com'
MAIL_PASSWORD = 'qaz135wsx'


def contracts_added_by(user_id):
    contracts = db.session.query(
                            CompaniesContract.title,
                            Contract.id,
                            Contract.created_at,
                            Contract.sent_at,
                            Contract.signed_at,
                            Contract.agreed_at,
                            Company.name)\
        .select_from(Contract)\
        .join(CompaniesContract)\
        .join(Company, Company.id == CompaniesContract.recipient_id)\
        .filter(Contract.created_by == user_id)\
        .all()
    contracts_list = []
    for c in contracts:
        contracts_list.append({
            'id': c.id,
            'title': c.title,
            'created_at': c.created_at,
            'sent_at': c.sent_at,
            'signed_at': c.signed_at,
            'agreed_at': c.agreed_at,
            'agreed_by_all': is_contract_agreed_by_all(c.id),
            'name': c.name,
        })
    return contracts_list


def contract_by_id(contract_id):
    # c = Contract.query.get(int(contract_id))
    c = db.session.query(
        CompaniesContract.title,
        Contract.id,
        Contract.created_at,
        Contract.created_by,
        Contract.should_agreed_by,
        Contract.should_signed_by,
        Contract.sent_at,
        Contract.signed_at,
        Contract.agreed_at,
        Company.name) \
        .select_from(Contract) \
        .join(CompaniesContract) \
        .join(Company, Company.id == CompaniesContract.recipient_id) \
        .filter(Contract.id == contract_id) \
        .first()
    if c is None:
        raise Exception('Договор не найден')
    return {
        'id': c.id,
        'title': c.title,
        'created_at': c.created_at,
        'should_agreed_by': c.should_agreed_by,
        'should_signed_by': c.should_signed_by,
        'created_by': c.created_by,
        'sent_at': c.sent_at,
        'signed_at': c.signed_at,
        'agreed_at': c.agreed_at,
        'agreed_by_all': is_contract_agreed_by_all(c.id),
        'name': c.name,
    }


def contract_versions(contract_id):
    return ContractVersion.query.filter(ContractVersion.id == contract_id).all()


def contract_agreements(contract_id):
    r = db.session.query(User.id,
                     User.name,
                     Role.name.label('role_name'),
                     ContractVersionsUsers.status,
                     ContractVersionsUsers.id.label('cvu_id'),
                     ContractVersion.created_at)\
        .select_from(ContractVersionsUsers)\
        .join(ContractVersion, ContractVersionsUsers.contract_version_id == ContractVersion.id)\
        .join(User, ContractVersionsUsers.user_id == User.id)\
        .join(Role, User.role_id == Role.id)\
        .filter(ContractVersion.contract_id == contract_id)\
        .order_by(ContractVersionsUsers.id)\
        .distinct(User.id).all()
    r_list = []
    for t in r:
        ucs_list = []
        if t.status is False:
            ucs = UsersComments.query.filter(UsersComments.versions_users_id == t.cvu_id)
            for uc in ucs:
                ucs_list.append({
                    'clause': uc.clause,
                    'original': uc.original,
                    'modified': uc.modified
                })
        r_list.append({
            'id': t.id,
            'name': t.name,
            'role_name': t.role_name,
            'status': t.status,
            'cvu_id': t.cvu_id,
            'comments': ucs_list
        })
    return r_list


def users_by_company(company_id, can_sign=None):
    if can_sign is None:
        users = db.session.query(User.id, User.name, Role.name.label('role_name'))\
            .select_from(User).join(Role).join(Company)\
            .filter(Company.id == company_id).all()
    else:
        users = db.session.query(User.id, User.name, Role.name.label('role_name')) \
            .select_from(User).join(Role).join(Company) \
            .filter(Company.id == company_id, Role.can_sign == can_sign).all()
    return users


def sign_by_company(company_id):
    signs = db.session.query(User.id, User.name, Role.name.label('role_name'))\
        .join(Role).join(Company)\
        .filter(Company.id == company_id, Role.can_sign == True)
    return signs.all()


UPLOAD_FOLDER = 'files'
FOLDER_PDF = 'pdf'

def save_to(folder, file, upload_id):
    os.makedirs(folder, exist_ok=True)
    name, ext = os.path.splitext (file.filename)
    save_path = os.path.join(folder, secure_filename(upload_id+ext))
    file.save(save_path)
    return save_path


def convert_to(folder, source, timeout=None):
    args = ['sudo ' + libreoffice_exec() + ' --headless --convert-to pdf --outdir pdf '+ source]

    process = subprocess.run(args, shell=True, timeout=timeout, capture_output=True)
    filename = re.search('-> (.*?) using filter', process.stdout.decode())

    if filename is None:
        raise LibreOfficeError(process.stdout.decode())
    return filename.group(1)

def libreoffice_exec():
    # TODO: Provide support for more platforms
    if sys.platform == 'darwin':
        return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    return 'libreoffice'

class LibreOfficeError(Exception):
    def __init__(self, output):
        self.output = output

def uploads_url(path):
    return path.replace(app.config['FOLDER_PDF'], '/uploads')

def upload_contract(file, created_by, should_agreed_by):
    upload_id = str(uuid4())
    source = save_to(os.path.join(app.config['FOLDER_PDF'], 'files'), file, upload_id)

    result = convert_to(os.path.join(app.config['FOLDER_PDF'], 'pdf'), source, timeout=15)

    doc = docx.Document(source)
    body = doc._body._body

    ct = ContractsType.query.all()

    title = 'Договор без названия'
    doc_type = 0
    recipient_id = 0

    for p in doc.paragraphs:
        heading = re.match(r'^Договор[а-яА-Я ]* № ?[\d]+$', p.text)
        if heading is not None:
            title = heading.group(0)
            for ct_item in ct:
                if ct_item.token in title:
                    doc_type = ct_item.id
            break
    if title == 'Договор без названия':
        flash('В документе не найдено название договора')
        raise Exception()
    if doc_type == 0:
        flash('Невозможно определить тип договора')
        raise Exception()

    initiator = db.session.query(Company.id, Company.inn) \
        .select_from(Company) \
        .join(Role) \
        .join(User) \
        .filter(User.id == created_by) \
        .first()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inn = re.findall(r"ИНН ?(\d+)", paragraph.text)
                    if len(inn) > 0 and initiator.inn != inn[0]:
                        recipient = Company.query.filter(Company.inn == inn[0]).first()
                        if recipient is None:
                            flash('Компания с ИНН ' + inn[0] + ' не найдена в инструменте')
                            raise Exception()
                        recipient_id = recipient.id
                        break

    cc = CompaniesContract(title=title, type=doc_type, initiator_id=initiator['id'], recipient_id=recipient_id)

    should_signed = db.session.query(User.id) \
        .select_from(Company) \
        .join(Role) \
        .join(User) \
        .filter(Company.id == initiator.id, Role.can_sign == True) \
        .first()

    c = Contract(created_by=created_by,
                 should_agreed_by=should_agreed_by,
                 should_signed_by=should_signed.id)
    cv = ContractVersion(filename=upload_id)
    c.versions.append(cv)

    ctu = ContractsTypesUsers.query.filter(ContractsTypesUsers.contract_type_id == doc_type,
                                           ContractsTypesUsers.company_id == initiator.id)

    for ctu_item in ctu:
        cvu = ContractVersionsUsers(user_id=ctu_item.user_id)
        cv.agreements.append(cvu)

    cc.contracts.append(c)

    db.session.add(cc)
    db.session.commit()

    for ctu_item in ctu:
        u = User.query.get(ctu_item.user_id)
        FROM = MAIL_USERNAME
        TO = 'gooddimkin@yandex.ru'

        msg = 'Уважаемый(ая) ' + u.name + ',\n\n Вам пришёл договор на согласование.\n Для его просмотра перейдите по ссылке: http://kpint.site/web/login\n\n С уважением,\n СогласовательПлюс'
        msg = MIMEText('\n {}'.format(msg).encode('utf-8'), _charset='utf-8')

        smtpObj = smtplib.SMTP_SSL(MAIL_SERVER, MAIL_PORT)
        smtpObj.ehlo()
        smtpObj.login(MAIL_USERNAME, MAIL_PASSWORD)

        smtpObj.sendmail(FROM, TO, 
                'Subject: СогласовательПлюс - договор на согласование. \n{}'.format(msg).encode('utf-8'))
        smtpObj.quit()

    return c.id


def add_comment(cvu_id, status, comments=None):
    cvu = ContractVersionsUsers.query.get(cvu_id)
    cvu.status = status
    cvu.status_added_at = datetime.datetime.now()
    if comments is not None:
        for comment in comments:
            uc = UsersComments(clause=comment['clause'],
                               original=comment['original'],
                               modified=comment['modified'])
            cvu.comments.append(uc)
    db.session.add(cvu)
    db.session.commit()

    FROM = MAIL_USERNAME
    TO = 'gooddimkin@yandex.ru'

    msg = 'Уважаемый(ая) Виноградов Рудольф Наумович,\n\n Статус согласования договора изменён.\n Для просмотра комментариев перейдите по ссылке: http://kpint.site/web/login\n\n С уважением,\n СогласовательПлюс'
    msg = MIMEText('\n {}'.format(msg).encode('utf-8'), _charset='utf-8')

    smtpObj = smtplib.SMTP_SSL(MAIL_SERVER, MAIL_PORT)
    smtpObj.ehlo()
    smtpObj.login(MAIL_USERNAME, MAIL_PASSWORD)

    smtpObj.sendmail(FROM, TO, 
            'Subject: СогласовательПлюс - договор на согласование. \n{}'.format(msg).encode('utf-8'))
    smtpObj.quit()


def contracts_for(u_id):
    cv = db.session.query(CompaniesContract.title,
                          Company.name,
                          ContractVersion.created_at,
                          ContractVersionsUsers.id,
                          ContractVersionsUsers.status) \
        .select_from(ContractVersionsUsers) \
        .join(ContractVersion) \
        .join(Contract) \
        .join(CompaniesContract) \
        .join(Company, Company.id == CompaniesContract.recipient_id) \
        .filter(ContractVersionsUsers.status == None, ContractVersionsUsers.user_id == u_id)
    cv = cv.all()
    cv_list = []
    for c in cv:
        cv_list.append({
            'cvu_id': c.id,
            'title': c.title,
            'name': c.name,
            'created_at': c.created_at
        })
    return cv_list

def update_contract(c_id, file, agreements):
    upload_id = str(uuid4())
    source = save_to(os.path.join(app.config['FOLDER_PDF'], 'files'), file, upload_id)

    result = convert_to(os.path.join(app.config['FOLDER_PDF'], 'pdf'), source, timeout=15)

    c = Contract.query.get(int(c_id))
    cv = ContractVersion(filename=upload_id)
    c.versions.append(cv)

    for user in agreements:
        cvu = ContractVersionsUsers(user_id=user.id)
        cv.agreements.append(cvu)

    db.session.add(c)
    db.session.commit()

    for user in agreements:
        u = User.query.get(user.id)
        FROM = MAIL_USERNAME
        TO = 'gooddimkin@yandex.ru'

        msg = 'Уважаемый(ая), ' + u.name + '\n\n Вам пришёл обновленный договор на согласование.\n Для его просмотра перейдите по ссылке: http://kpint.site/web/login\n\n С уважением,\n СогласовательПлюс'
        msg = MIMEText('\n {}'.format(msg).encode('utf-8'), _charset='utf-8')

        smtpObj = smtplib.SMTP_SSL(MAIL_SERVER, MAIL_PORT)
        smtpObj.ehlo()
        smtpObj.login(MAIL_USERNAME, MAIL_PASSWORD)

        smtpObj.sendmail(FROM, TO, 
                'Subject: СогласовательПлюс - договор на согласование. \n{}'.format(msg).encode('utf-8'))
        smtpObj.quit()

    return c.id


def users_by_contract(c_id):
    # TODO: change query: get only False for current version of doc
    r = db.session.query(User.id,
                         User.name,
                         Role.name.label('role_name'),
                         ContractVersionsUsers.status,
                         ContractVersionsUsers.id.label('cvu_id')) \
        .select_from(ContractVersionsUsers) \
        .join(ContractVersion, ContractVersionsUsers.contract_version_id == ContractVersion.id) \
        .join(User, ContractVersionsUsers.user_id == User.id) \
        .join(Role, User.role_id == Role.id) \
        .filter(ContractVersion.contract_id == c_id, ContractVersionsUsers.status == False) \
        .order_by(ContractVersionsUsers.id) \
        .distinct(User.id).all()
    return r


def is_contract_still_agrees(c_id):
    sql = """SELECT COUNT(*)
                FROM contract_versions_users cvu
                WHERE cvu.contract_version_id ==
                (SELECT cv.id
                FROM contract c
                JOIN contract_version cv on c.id = cv.contract_id
                WHERE c.id = :c_id
                ORDER BY cv.id DESC
                LIMIT 1)
                AND cvu.status IS NULL"""

    result = db.session.execute(sql, {'c_id': c_id})
    count = result.scalar()

    return count > 0


def is_contract_agreed_by_all(c_id):
    sql = """SELECT COUNT(*)
                FROM contract_versions_users cvu
                WHERE cvu.contract_version_id ==
                (SELECT cv.id
                FROM contract c
                JOIN contract_version cv on c.id = cv.contract_id
                WHERE c.id = :c_id
                ORDER BY cv.id DESC
                LIMIT 1)
                AND (cvu.status == 0 OR cvu.status IS NULL)"""

    result = db.session.execute(sql, {'c_id': c_id})
    count = result.scalar()

    return not count > 0
