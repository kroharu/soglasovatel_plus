import re

import docx

doc = docx.Document('Договор аренды нежилого помещения (1).docx')

body = doc._body._body

# print(body)

types = {
    'аренд': 'аренда',
    'постав': 'поставка'
}

# print(body.xml)

# for p in doc.paragraphs:
#     heading = re.match(r'^Договор[а-яА-Я ]* № ?[\d]+$', p.text)
#     if heading is not None:
#         text = heading.group(0)
#         print(text)
#         for t in types:
#             if t in text:
#                 print(types[t])
#         break

inns = []

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                inn = re.findall(r"ИНН ?(\d+)", paragraph.text)
                if len(inn) > 0:
                    print(inn[0])
                    inns.append(inn[0])

print(inns)
